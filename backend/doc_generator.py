"""
doc_generator.py
-----------------
Converts the Markdown study-notes string produced by ai.py into a fully
styled Microsoft Word (.docx) document using python-docx, and saves it to a
temporary directory with a unique, collision-free filename.

Public API:
    generate_docx(markdown_text: str, video_title: str = "NoteTube AI Notes") -> str
        Returns the absolute filepath of the generated .docx file.
"""

import os
import re
import uuid

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# Style palette — matches the frontend's dark/amber "NoteTube AI" identity,
# translated into a professional, print-friendly corporate scheme.
# ---------------------------------------------------------------------------

COLOR_HEADING_1 = RGBColor(0x1B, 0x24, 0x38)   # deep slate navy
COLOR_HEADING_2 = RGBColor(0x2E, 0x3D, 0x5C)   # slate blue
COLOR_HEADING_3 = RGBColor(0x9A, 0x6A, 0x1F)   # muted amber/gold
COLOR_ACCENT = RGBColor(0xC9, 0x8A, 0x2D)      # amber accent
COLOR_BODY = RGBColor(0x22, 0x22, 0x22)
COLOR_TABLE_HEADER_BG = "1B2438"

FONT_BODY = "Calibri"
FONT_HEADING = "Arial"

TEMP_DIR_DEFAULT = "temp_files"


def _ensure_temp_dir(temp_dir: str) -> str:
    os.makedirs(temp_dir, exist_ok=True)
    return temp_dir


def _set_cell_background(cell, hex_color: str):
    shading_elm = cell._tc.get_or_add_tcPr()
    shd = shading_elm.makeelement(
        qn("w:shd"), {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color}
    )
    shading_elm.append(shd)


def _configure_base_styles(document: Document):
    normal = document.styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(11)
    normal.font.color.rgb = COLOR_BODY
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for section in document.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)


def _add_title_page(document: Document, title: str):
    title_para = document.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.font.name = FONT_HEADING
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = COLOR_HEADING_1

    subtitle_para = document.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_para.add_run("Generated Study Notes — NoteTube AI")
    sub_run.font.name = FONT_BODY
    sub_run.font.size = Pt(12)
    sub_run.font.italic = True
    sub_run.font.color.rgb = COLOR_ACCENT

    document.add_paragraph()  # spacer
    document.add_page_break()


def _add_heading(document: Document, text: str, level: int):
    heading = document.add_heading(level=min(level, 3))
    heading.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    heading.paragraph_format.space_after = Pt(8)
    run = heading.add_run(text)
    run.font.name = FONT_HEADING
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = COLOR_HEADING_1
    elif level == 2:
        run.font.size = Pt(16)
        run.font.color.rgb = COLOR_HEADING_2
    else:
        run.font.size = Pt(13)
        run.font.color.rgb = COLOR_HEADING_3


_INLINE_BOLD_PATTERN = re.compile(r"\*\*(.+?)\*\*")
_INLINE_CODE_PATTERN = re.compile(r"`([^`]+)`")


def _add_paragraph_with_inline_formatting(document: Document, text: str, bullet: bool = False, numbered: bool = False):
    style = None
    if bullet:
        style = "List Bullet"
    elif numbered:
        style = "List Number"

    paragraph = document.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(6)

    # Tokenize on bold (**text**) and inline code (`text`) markers, preserving order.
    token_pattern = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")
    tokens = token_pattern.split(text)

    for token in tokens:
        if not token:
            continue
        bold_match = _INLINE_BOLD_PATTERN.fullmatch(token)
        code_match = _INLINE_CODE_PATTERN.fullmatch(token)
        if bold_match:
            run = paragraph.add_run(bold_match.group(1))
            run.bold = True
            run.font.color.rgb = COLOR_HEADING_1
        elif code_match:
            run = paragraph.add_run(code_match.group(1))
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x8A, 0x4B, 0x08)
        else:
            paragraph.add_run(token)

    return paragraph


def _add_code_block(document: Document, code_lines: list):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.left_indent = Inches(0.3)
    run = paragraph.add_run("\n".join(code_lines))
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xEE, 0xEE, 0xEE)
    # Simulate a code block background via shading on the paragraph
    shading_elm = paragraph._p.get_or_add_pPr()
    shd = shading_elm.makeelement(
        qn("w:shd"), {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): "1B2438"}
    )
    shading_elm.append(shd)


def _add_table(document: Document, header_row: list, data_rows: list):
    table = document.add_table(rows=1, cols=len(header_row))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_cells = table.rows[0].cells
    for idx, header_text in enumerate(header_row):
        header_cells[idx].text = ""
        run = header_cells[idx].paragraphs[0].add_run(header_text.strip())
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = FONT_HEADING
        run.font.size = Pt(10.5)
        _set_cell_background(header_cells[idx], COLOR_TABLE_HEADER_BG)

    for row_data in data_rows:
        row_cells = table.add_row().cells
        for idx, cell_text in enumerate(row_data):
            if idx < len(row_cells):
                row_cells[idx].text = ""
                run = row_cells[idx].paragraphs[0].add_run(cell_text.strip())
                run.font.size = Pt(10.5)
                run.font.name = FONT_BODY

    document.add_paragraph().paragraph_format.space_after = Pt(6)


def _parse_markdown_to_docx(document: Document, markdown_text: str):
    lines = markdown_text.split("\n")
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Fenced code block
        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            if code_lines:
                _add_code_block(document, code_lines)
            continue

        # Markdown table (header row + separator row like |---|---|)
        if stripped.startswith("|") and i + 1 < n and re.match(r"^\|?[\s:|-]+\|?$", lines[i + 1].strip()):
            header_cells = [c.strip() for c in stripped.strip("|").split("|")]
            i += 2  # skip header + separator
            data_rows = []
            while i < n and lines[i].strip().startswith("|"):
                row_cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                data_rows.append(row_cells)
                i += 1
            _add_table(document, header_cells, data_rows)
            continue

        # Headings
        heading_match = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            _add_heading(document, text, level)
            i += 1
            continue

        # Bullet list item
        bullet_match = re.match(r"^[-*]\s+(.*)$", stripped)
        if bullet_match:
            _add_paragraph_with_inline_formatting(document, bullet_match.group(1), bullet=True)
            i += 1
            continue

        # Numbered list item
        numbered_match = re.match(r"^\d+[.)]\s+(.*)$", stripped)
        if numbered_match:
            _add_paragraph_with_inline_formatting(document, numbered_match.group(1), numbered=True)
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^-{3,}$", stripped):
            i += 1
            continue

        # Default: regular paragraph
        _add_paragraph_with_inline_formatting(document, stripped)
        i += 1


def generate_docx(markdown_text: str, video_title: str = "NoteTube AI Notes", temp_dir: str = TEMP_DIR_DEFAULT) -> str:
    """
    Render the given Markdown study notes into a styled .docx file.

    Args:
        markdown_text: the Markdown document produced by ai.generate_notes().
        video_title: title to render on the cover page and use in the filename.
        temp_dir: directory (relative or absolute) to save the file into.

    Returns:
        str: absolute filepath of the generated .docx file.
    """
    if not markdown_text or not markdown_text.strip():
        raise ValueError("Cannot generate a document from empty Markdown content.")

    temp_dir = _ensure_temp_dir(temp_dir)

    document = Document()
    _configure_base_styles(document)
    _add_title_page(document, video_title)
    _parse_markdown_to_docx(document, markdown_text)

    safe_title = re.sub(r"[^a-zA-Z0-9_-]+", "_", video_title.strip())[:40] or "notes"
    unique_id = uuid.uuid4().hex[:10]
    filename = f"{safe_title}_{unique_id}.docx"
    filepath = os.path.abspath(os.path.join(temp_dir, filename))

    document.save(filepath)
    return filepath
